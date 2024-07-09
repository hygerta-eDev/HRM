from fastapi import APIRouter, Depends, HTTPException, Form,Query
from sqlalchemy.orm import Session
from .employeeService import EmployeeService
from Schema.employeeSchema import EmployeeCreate,EmployeeUpdate,UsernameRequest,UsernameResponse,documentEmployee
from Config.database import get_db
from typing import List,Dict
from Schema.enums.Marital_status import MaritalStatus
from Schema.enums.genders import Genders
from Schema.enums.city import cities_data,CityResponse,City
from docx import Document
from docx.shared import Pt
import os
from fastapi.responses import FileResponse
from pathlib import Path
import aiofiles
from Schema.DMS import DocumentBase
from Models.employeeModel import Employees
# from Models.DMS import DocumentS
from datetime import datetime
import Models
router = APIRouter(prefix="/employees", tags=["Employee"])

@router.get("/")
def get_all_employees(db: Session = Depends(get_db)):
    return EmployeeService.get_all_employees(db=db)

@router.get("/employees/{employee_id}")
def get_employee(employee_id: int, db: Session = Depends(get_db)):
    employee = EmployeeService.get_employee_by_id(db, employee_id)
    if employee is None:
        raise HTTPException(status_code=404, detail="Employee not found")
    return employee

@router.post("/create_employee")
def create_employee(Employee: EmployeeCreate, db: Session = Depends(get_db)):
    return EmployeeService.create_employee(Employee, db)


# @router.put("/{employee_id}")
# def update_appointment(employee_id: int, Employee: EmployeeUpdate, db: Session = Depends(get_db)):
#     return EmployeeService.update_appointment(employee_id=employee_id, Employee=Employee, db=db)
# @router.put("/employees/{employee_id}/update_appointment")
# def update_employee_appointment(employee_id: int, employee_update: EmployeeUpdate, db: Session = Depends(get_db)):
#     db_employee = EmployeeService.get_employee_by_id(db, employee_id)
#     if db_employee is None:
#         raise HTTPException(status_code=404, detail="Employee not found")
    
#     # Update only the fields that are provided in the request
#     for field, value in employee_update.dict().items():
#         if value is not None:
#             setattr(db_employee, field, value)
    
#     db.commit()
#     db.refresh(db_employee)
    
#     return db_employee
@router.put("/update_employee/{employee_id}")
def update_employee(employee_id: int, Employee: EmployeeUpdate, db: Session = Depends(get_db)):
    return EmployeeService.update_employee(employee_id=employee_id, Employee=Employee, db=db)
@router.delete("/delete_employee/{employee_id}")
def delete_employee(employee_id: int, db: Session = Depends(get_db)):
    return EmployeeService.delete_employee(employee_id=employee_id, db=db)

# @router.get("/marital_status", response_model=List[MaritalStatus])
# async def get_marital_status_options():
#     return [status.to_dict() for status in MaritalStatus]
@router.get("/marital_status", response_model=List[MaritalStatus])
async def read_marital_status():
    return list(MaritalStatus)

# @router.get("/genders", response_model=List[str])
# def get_genders_options():
#     genders_options = [status.value for status in Genders]
#     return genders_options

@router.get("/genders", response_model=List[Genders])
async def read_genders():
    return list(Genders)

@router.get("/check-number-exists")
def check_number_exists(number: str = Query(...), db: Session = Depends(get_db)):
    existing_employee = EmployeeService.get_employee_by_number(db, number)
    if existing_employee:
        return {"exists": True}
    else:
        return {"exists": False}

@router.get("/check-personal-number-exists")
def check_personal_number_exists(personal_number: str = Query(...), db: Session = Depends(get_db)):
    existing_employee = EmployeeService.get_employee_by_personal_number(db, personal_number)
    if existing_employee:
        return {"exists": True}
    else:
        return {"exists": False}
    

@router.get("/download/cv/{employee_id}")
def download_cv(employee_id: int, db: Session = Depends(get_db)):
    try:
        # Call the download_cv method of the service class
        return EmployeeService.download_cv(employee_id,db)
    except KeyError:
        raise HTTPException(status_code=404, detail="Employee not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    


@router.get("/last_employee_id")
def get_last_employee_id(db: Session = Depends(get_db)):
    return EmployeeService.get_last_employee_id(db=db)


@router.get("/cities", response_model=List[CityResponse])
async def get_cities():
    city_responses = [
        CityResponse(city_name=city_name, zip_codes=cities_data[city_name]) 
        for city_name in cities_data.keys()
    ]
    return city_responses

@router.get("/zipcodes", response_model=CityResponse)
def get_cities(city_name: str):
    if city_name not in cities_data:
        raise HTTPException(status_code=404, detail="City not found")
    zip_codes = cities_data[city_name]
    return CityResponse(city_name=city_name, zip_codes=zip_codes)

@router.get("/generate_unique_number")
def generate_unique_number(db: Session = Depends(get_db)):
    unique_number = EmployeeService.generate_unique_number(db)
    return {"number": unique_number}

@router.post("/generate-username", response_model=UsernameResponse)
def generate_username(request_data: UsernameRequest):
    if request_data.name and request_data.last_name:
        username = f"{request_data.name.lower()}.{request_data.last_name.lower()}"
        return {"username": username}
    else:
        return {"username": None}
    


# C:\Users\hygerta.hulaj\Desktop\HRM\HRM\HRM_backend\DMS\F-057 Kontrata e punes E-Tech.docx
# def generate_document(employee: documentEmployee, template_path: str, output_path: str):
#     doc = Document(template_path)

#     def replace_placeholder(run, placeholder, value):
#         if placeholder in run.text:
#             run.text = run.text.replace(placeholder, str(value))

#     placeholders = {
#         "{{name}}": employee.name,
#         "{{personal_number}}": employee.personal_number,
#         "{{date_of_birth}}": employee.date_of_birth.strftime("%Y-%m-%d"),
#         "{{place_of_birth}}": employee.city,
#         "{{address}}": employee.street,
#         "{{job_position}}": employee.job_position_id,
#         "{{date_hired}}": employee.date_hired.strftime("%Y-%m-%d"),
#         "{{contract_end_date}}": employee.contract_end_date.strftime("%Y-%m-%d") if employee.contract_end_date else "",
#         "{{today}}": employee.today.strftime("%Y-%m-%d"),
#         "{{salary}}": employee.salary,
#     }

#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             for placeholder, value in placeholders.items():
#                 replace_placeholder(run, placeholder, value)

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     for run in paragraph.runs:
#                         for placeholder, value in placeholders.items():
#                             replace_placeholder(run, placeholder, value)

#     doc.save(output_path)


def generate_document(employee: documentEmployee, template_path: str, output_path: str) -> str:
    doc = Document(template_path)
    # def replace_placeholder(run, placeholder, value):
    #     if placeholder in run.text:
    #         run.text = run.text.replace(placeholder, str(value))
    def replace_placeholder(paragraph, placeholder, value):
        # Concatenate all text within the paragraph
        text = ''.join(run.text for run in paragraph.runs)
        if placeholder in text:
            # Replace placeholder in the concatenated text
            text = text.replace(placeholder, str(value))
            # text.bold = True

            # Clear existing runs and add the replaced text as a new run
            for run in paragraph.runs:
                run.text = ''
            new_run = paragraph.add_run(text)
            # new_run.font.size = Pt(12)
            # new_run.bold = True

    placeholders = {
        "{{name}}": employee.name,
        "{{personal_number}}": employee.personal_number,
        "{{date_of_birth}}": employee.date_of_birth.strftime("%Y-%m-%d"),
        "{{place_of_birth}}": employee.city,
        "{{address}}": employee.street,
        "{{job_position}}": employee.job_position_id,
        "{{date_hired}}": employee.date_hired.strftime("%Y-%m-%d"),
        "{{contract_end_date}}": employee.contract_end_date.strftime("%Y-%m-%d") if employee.contract_end_date else "",
        "{{today}}": employee.created_at.strftime("%Y-%m-%d"),
        "{{salary}}": employee.salary,
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            replace_placeholder(paragraph, placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        replace_placeholder(paragraph, placeholder, value)

    doc.save(output_path)
    return output_path

# @router.post("/generate-document")
# def generate_document_endpoint(employee: EmployeeBase):
#     template_path = "./DMS/F-057 Kontrata e punes E-Tech.docx"  # Path to your template file
#     output_filename = f"{employee.name}_contract.docx"
#     output_path = os.path.join("./DMS/", output_filename)

#     if not os.path.exists(template_path):
#         raise HTTPException(status_code=404, detail=f"Template not found at '{template_path}'")

#     generate_document(employee, template_path, output_path)
#     return {"detail": "Document generated successfully", "filename": output_filename}



@router.post("/generate-document")
def generate_document_endpoint(employee: documentEmployee):
    template_path = "./DMS/F-057 Kontrata e Punës - eDev.docx" 
    output_filename = f"{employee.name}_contract.docx"
    output_path = os.path.join("./DMS/contract/", output_filename)

    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Template not found at '{template_path}'")

    generate_document(employee, template_path, output_path)

    # Ensure the file exists before returning
    if os.path.exists(output_path):
        return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=output_filename)
    else:
        raise HTTPException(status_code=404, detail=f"Generated document not found")
@router.get("/download-document")
def download_document():
    output_path = "./DMS/output.docx"
    
    if not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="Generated document not found")



@router.post("/generate_and_attach_document/{employee_id}")
async def generate_and_attach_document(employee_id: int, employee: documentEmployee, db: Session = Depends(get_db)):
    # Fetch employee details from the database
    employee_record = db.query(Employees).filter(Employees.id == employee_id).first()
    if not employee_record:
        raise HTTPException(status_code=404, detail="Employee not found")

    # Print current working directory for debugging
    print(f"Current working directory: {os.getcwd()}")

    # Define template and output paths
    template_path = Path("./DMS/F-057 Kontrata e Punës - eDev.docx").resolve()
    output_path = Path(f"./DMS/contract/{employee_record.name}_contract.docx").resolve()

    # Check if the template file exists
    if not template_path.exists():
        raise HTTPException(status_code=404, detail=f"Template not found at '{template_path}'")

    # Generate the document
    document_path = generate_document(employee, template_path, output_path)
    print(document_path)
    # Ensure the document was generated and the path is not None
    # if not document_path or not Path(document_path).exists():
    #     raise HTTPException(status_code=500, detail="Document generation failed")

    # Save the generated document in the database
    document = Models.DMS.Document(
        title=f"{employee.name}_contract.docx",
        description="Kontratat e Punësimit",
        file_path=str(document_path),
        employee_id=employee_id,
        category_id=2,  
        created_at=datetime.now(),
        updated_at=datetime.now()
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return document

@router.get("/employee/{employee_id}/total_leave_quota_amount")
def read_total_leave_quota_amount(employee_id: int, db: Session = Depends(get_db)):
    total_amount = EmployeeService.get_total_leave_quota_amount(employee_id, db)
    if total_amount is None:
        raise HTTPException(status_code=404, detail="Employee not found")
    return {"employee_id": employee_id, "total_leave_quota_amount": total_amount}